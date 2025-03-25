"""
This is about converting docx or doc to .txt
"""
import os
from docx import Document

class DocConvert:
    """_summary_
    """
    def __init__(self, document_path, txt_path):
        self.document_path = document_path
        self.txt_path = txt_path


    def docx_to_txt(self, docx_path, txt_path):
        """_summary_

        Args:
            docx_path (_type_): _description_
            txt_path (_type_): _description_
        """
        output_dir = os.path.dirname(txt_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        document = Document(docx_path)
        with open(txt_path, 'w', encoding="utf-8") as f:
            for p in document.paragraphs:
                f.write(p.text + '\n')
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    f.write("\t".join(row_text) + "\n")
